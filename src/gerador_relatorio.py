from docx import Document

from docx.shared import Pt

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import re


# =========================================================
# LIMPAR TEXTO
# =========================================================

def limpar_texto_xml(texto):

    if texto is None:
        return ""

    texto = str(texto)

    # Remove caracteres incompatíveis com Word/XML
    texto = re.sub(
        r"[\x00-\x08\x0B\x0C\x0E-\x1F]",
        "",
        texto
    )

    # Corrige quebras de linha vindas do PDF
    texto = texto.replace("\r", "\n")
    texto = re.sub(r"-\n", "", texto)
    texto = re.sub(r"\n+", " ", texto)

    # Remove espaços exagerados
    texto = re.sub(r"\s{2,}", " ", texto)

    return texto.strip()


# =========================================================
# FORMATAÇÃO ABNT
# =========================================================

def formatar_paragrafo(paragrafo):

    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    paragrafo.paragraph_format.first_line_indent = Pt(35)

    paragrafo.paragraph_format.line_spacing = 1.5

    for run in paragrafo.runs:

        run.font.name = "Times New Roman"

        run.font.size = Pt(12)


# =========================================================
# ADICIONAR TEXTO FORMATADO
# =========================================================

def adicionar_texto(documento, texto):

    p = documento.add_paragraph(
        limpar_texto_xml(texto)
    )

    formatar_paragrafo(p)

    return p


# =========================================================
# GERAR RELATÓRIO
# =========================================================

def gerar_relatorio(resultados, caminho_saida):

    documento = Document()


    # =====================================================
    # TÍTULO
    # =====================================================

    titulo = documento.add_heading(
        "ENAMED INTELLIGENCE",
        level=1
    )

    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    subtitulo = documento.add_heading(
        "Relatório Inteligente de Mapeamento de Questões",
        level=2
    )

    subtitulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    introducao = """
Este relatório apresenta as questões extraídas automaticamente da prova selecionada,
integrando inteligência artificial, análise curricular, identificação de competências,
habilidades e conteúdos médicos compatíveis com os planos de aula cadastrados no sistema.
"""

    adicionar_texto(
        documento,
        introducao
    )


    # =====================================================
    # QUESTÕES
    # =====================================================

    for resultado in resultados:

        cabecalho = documento.add_heading(
            f"QUESTÃO {resultado['numero']}",
            level=2
        )

        cabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        adicionar_texto(
            documento,
            f"Fonte: {resultado['fonte']}"
        )

        adicionar_texto(
            documento,
            f"Ano: {resultado['ano']}"
        )

        adicionar_texto(
            documento,
            f"Arquivo de origem: {resultado['arquivo']}"
        )

        adicionar_texto(
            documento,
            f"Disciplina recomendada: {resultado['disciplina']}"
        )

        adicionar_texto(
            documento,
            f"Compatibilidade: {resultado['score']}%"
        )

        adicionar_texto(
            documento,
            f"Conteúdo provável: {resultado['conteudo']}"
        )


        # ================================================
        # ANÁLISE IA
        # ================================================

        subtitulo_ia = documento.add_heading(
            "Análise Pedagógica da IA",
            level=3
        )

        subtitulo_ia.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        adicionar_texto(
            documento,
            resultado['analise_ia']
        )


        # ================================================
        # QUESTÃO COMPLETA
        # ================================================

        subtitulo_questao = documento.add_heading(
            "Texto Completo da Questão",
            level=3
        )

        subtitulo_questao.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


        adicionar_texto(
            documento,
            resultado['questao']
        )

        documento.add_page_break()


    # =====================================================
    # SALVAR
    # =====================================================

    documento.save(caminho_saida)